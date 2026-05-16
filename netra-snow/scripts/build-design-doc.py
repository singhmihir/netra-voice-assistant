"""Build the Netra Technical Design Document (.docx).

Run:  python netra-snow/scripts/build-design-doc.py
Out:  netra-snow/docs/Netra-Technical-Design.docx
"""
from __future__ import annotations
import datetime
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
OUT_DIR = os.path.join(ROOT, "docs")
DIAG_DIR = os.path.join(OUT_DIR, "diagrams")
os.makedirs(OUT_DIR, exist_ok=True)
OUT = os.path.join(OUT_DIR, "Netra-Technical-Design.docx")

def add_image(doc, name, width_in=6.5, caption=None):
    """Insert an image from docs/diagrams/<name>.png with an optional caption."""
    path = os.path.join(DIAG_DIR, name + ".png")
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(path, width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = TEXT_GREY

BRAND_BLUE  = RGBColor(0x1A, 0x56, 0xDB)
BRAND_DARK  = RGBColor(0x0E, 0x2E, 0x7A)
TEXT_GREY   = RGBColor(0x47, 0x55, 0x69)
LIGHT_GREY  = RGBColor(0xE2, 0xE8, 0xF0)

# ---------- helpers --------------------------------------------------------
def shade(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_h1(doc, text):
    p = doc.add_paragraph(text, style='Heading 1')
    return p

def add_h2(doc, text):
    return doc.add_paragraph(text, style='Heading 2')

def add_h3(doc, text):
    return doc.add_paragraph(text, style='Heading 3')

def add_p(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(text, style='List Number')
    return p

def add_code(doc, text):
    """Monospaced code block in a single-cell table with grey shading."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    shade(cell, "F1F4F9")
    # remove default paragraph
    cell.text = ""
    for line in text.splitlines():
        p = cell.add_paragraph()
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        # paragraph spacing tight
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    add_p(doc, "")

def add_table(doc, headers, rows, col_widths_in=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    tbl.autofit = False
    if col_widths_in:
        for i, w in enumerate(col_widths_in):
            for cell in tbl.columns[i].cells:
                cell.width = Inches(w)
    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell, "1A56DB")
    # Body rows
    for ri, row_data in enumerate(rows, start=1):
        for ci, val in enumerate(row_data):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.add_run(str(val))
    add_p(doc, "")
    return tbl

def page_break(doc):
    doc.add_page_break()

# ---------- document setup ------------------------------------------------
doc = Document()

# Set default font and page size
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Heading style tweaks
for hname, size, color in [('Heading 1', 22, BRAND_DARK),
                             ('Heading 2', 16, BRAND_BLUE),
                             ('Heading 3', 13, BRAND_DARK)]:
    s = doc.styles[hname]
    s.font.name = 'Calibri'
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = color

# ============================================================
# Cover page
# ============================================================
cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_title.add_run('Netra Voice Assistant')
r.font.size = Pt(40); r.bold = True; r.font.color.rgb = BRAND_DARK

cover_subtitle = doc.add_paragraph()
cover_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_subtitle.add_run('Technical Design Document')
r.font.size = Pt(20); r.italic = True; r.font.color.rgb = BRAND_BLUE

add_p(doc, "")

# Tagline box (single-cell table)
t = doc.add_table(rows=1, cols=1)
t.autofit = False
t.columns[0].width = Inches(6.5)
c = t.cell(0, 0); c.width = Inches(6.5)
shade(c, "EFF6FF")
c.text = ""
para = c.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = para.add_run("An always-listening, voice-first assistant for blind and visually-impaired\n"
                 "ServiceNow users. Native to the platform. No external services.")
r.italic = True; r.font.size = Pt(12); r.font.color.rgb = TEXT_GREY

add_p(doc, "")
add_p(doc, "")
add_p(doc, "")

# Meta block (version, date, scope, author)
meta = doc.add_table(rows=4, cols=2)
meta.autofit = False
for col, w in [(0, 2.0), (1, 4.5)]:
    for cell in meta.columns[col].cells:
        cell.width = Inches(w)
for r_idx, (label, value) in enumerate([
    ("Version",       "3.0.0"),
    ("Date",          datetime.date.today().strftime("%d %B %Y")),
    ("Application",   "netra  (scope auto-generated as x_<companyId>_netra)"),
    ("Document type", "Technical Design / Architecture"),
]):
    cell_l = meta.rows[r_idx].cells[0]
    cell_l.text = ""
    p = cell_l.paragraphs[0]; rr = p.add_run(label); rr.bold = True; rr.font.color.rgb = TEXT_GREY
    cell_r = meta.rows[r_idx].cells[1]
    cell_r.text = ""
    cell_r.paragraphs[0].add_run(value)

page_break(doc)

# ============================================================
# Table of contents (manual — Word can refresh on open)
# ============================================================
add_h1(doc, "Contents")
for item in [
    ("1. Executive summary", 3),
    ("2. Goals & non-goals", 4),
    ("3. System architecture", 5),
    ("4. Component reference", 7),
    ("5. Data model", 10),
    ("6. Voice interaction flow", 12),
    ("7. Voice command grammar", 14),
    ("8. Accessibility design", 16),
    ("9. Security & data residency", 17),
    ("10. Installation guide", 18),
    ("11. Operational concerns", 20),
    ("12. Test plan", 21),
    ("13. Roadmap & future work", 22),
    ("Appendix A — Source file index", 23),
]:
    p = doc.add_paragraph()
    r = p.add_run(item[0]); r.font.size = Pt(11)
    r2 = p.add_run("\t" + str(item[1])); r2.font.size = Pt(11)

page_break(doc)

# ============================================================
# 1. Executive summary
# ============================================================
add_h1(doc, "1. Executive summary")
add_p(doc, "Netra is a fully-native ServiceNow scoped application that turns the platform into a voice-first assistant for blind and visually-impaired users. It listens for the wake word \"Netra\" inside any open Service Portal tab, transcribes the user's commands, and uses a deterministic, regex-based intent parser plus GlideRecord operations to create, read, resolve and update tickets — and also surface approvals, search knowledge articles, read chat messages, and announce new assignments proactively.")
add_p(doc, "Every business operation happens inside the customer's ServiceNow instance. No tokens are exchanged with external APIs and no message text leaves the platform. The only external interaction is the browser-native Web Speech API for audio → text conversion (handled by Chrome / Edge / Safari themselves, not by Netra), and a text-only mode is provided for users who require strict in-instance processing.", italic=False)
add_p(doc, "")
add_p(doc, "Key facts:", bold=True)
add_bullet(doc, "Single deployable Update Set XML (~150 KB) containing one Fix Script.")
add_bullet(doc, "Scope-agnostic: detects the scope auto-generated for an app named 'netra'.")
add_bullet(doc, "9 Script Includes, 4 tables, 1 Business Rule, 1 Scheduled Job (every 3 min), 1 Scripted REST API with 2 resources, 1 Service Portal widget.")
add_bullet(doc, "Zero external dependencies, zero API keys, zero recurring cost.")
add_bullet(doc, "Re-running the Fix Script is idempotent — safe to re-deploy.")

page_break(doc)

# ============================================================
# 2. Goals & non-goals
# ============================================================
add_h1(doc, "2. Goals & non-goals")

add_h2(doc, "2.1 Goals")
add_bullet(doc, "Voice-first ticket management — create, list, resolve, update, status, search.")
add_bullet(doc, "Always-on wake word inside the Service Portal tab.")
add_bullet(doc, "Proactive notifications when something needs the user's attention.")
add_bullet(doc, "Conversational pause / resume so the user can control interruption volume.")
add_bullet(doc, "Knowledge-base search and approval handling.")
add_bullet(doc, "Multi-turn dialogues (e.g. \"open a ticket\" → \"what's the issue?\" → user replies).")
add_bullet(doc, "WCAG 2.1 AA accessibility — keyboard, screen-reader, high contrast.")
add_bullet(doc, "Single-file deployment via Update Set XML.")

add_h2(doc, "2.2 Non-goals (v3)")
add_bullet(doc, "Cross-instance federation or external chat integrations (Teams, Slack).")
add_bullet(doc, "Custom on-device speech-to-text — relies on browser-provided Web Speech API.")
add_bullet(doc, "Voice biometrics or speaker identification.")
add_bullet(doc, "Translation / multi-language NLU (English-only intent parsing).")
add_bullet(doc, "Mobile native app — Service Portal works on mobile browsers, but no separate iOS / Android app.")

page_break(doc)

# ============================================================
# 3. System architecture
# ============================================================
add_h1(doc, "3. System architecture")
add_p(doc, "Netra is a self-contained scoped application. Every artifact lives inside the customer's instance.")

add_h2(doc, "3.1 Logical architecture")
add_image(doc, "architecture", width_in=6.5, caption="Figure 1 — Netra v3 logical architecture. Every component lives inside the customer's ServiceNow instance.")
add_p(doc, "ASCII representation of the same diagram (for accessibility and machine-readable transcripts):", italic=True)
add_code(doc, r"""
+--------------------------------------------------------------------------+
|                       Browser (Chrome / Edge / Safari)                   |
|                                                                          |
|   +--------------------------------------------------------------+      |
|   |  Service Portal page  (any container hosts the netra-mic     |      |
|   |                        widget; floats fixed bottom-right)    |      |
|   |                                                              |      |
|   |   Wake-word listener  --->  Web Speech STT  --->             |      |
|   |   (continuous)              (browser-native)                 |      |
|   |                                                              |      |
|   |   <---  Speech Synthesis  <---  Netra reply text             |      |
|   |        (Indian-English voice preferred)                      |      |
|   +-------------------------------|------------------------------+      |
|                                   | POST /api/<scope>/voice/command     |
|                                   | GET  /api/<scope>/voice/notifications|
+-----------------------------------|--------------------------------------+
                                    |
+-----------------------------------v--------------------------------------+
|                       ServiceNow Instance  (scope x_<id>_netra)          |
|                                                                          |
|   Scripted REST API                                                      |
|        |                                                                 |
|        v                                                                 |
|   NetraIntent  ---> NetraResponder  --->  NetraTools  (GlideRecord)      |
|                              \                                           |
|                               +-->  NetraKnowledge  (kb_knowledge)       |
|                               +-->  NetraChat       (sys_cs_*)           |
|                               +-->  NetraContext    (focus memory)       |
|                               +-->  NetraSummarizer (extractive)         |
|                               +-->  NetraNavigator  (portal URL maps)    |
|                                                                          |
|   Business Rule on sys_journal_field   ----+                             |
|   Scheduled Job "Netra Watch" (3 min)  ----+--> netra_notification queue |
|        |                                        ^                        |
|        +--> NetraScanner --> incident/change/   |                        |
|                              sc_task/approvals  |                        |
|                                                  |                        |
|   netra_user_pref  (per-user state, pause)      |                        |
|   netra_context    (focused ticket per user)    |                        |
|   netra_watchlist  (per-user watched records)   |                        |
+---------------------------------------------------+--------------------+
                                                    | polled every 8s by  |
                                                    | the widget          |
                                                    +---------------------+
""")

add_h2(doc, "3.2 Process flow — voice command")
add_image(doc, "sequence-command", width_in=6.5, caption="Figure 2 — Sequence diagram for a voice command round-trip.")
add_p(doc, "The end-to-end flow for a typical command (\"resolve INC0001234\"):")
add_numbered(doc, "Widget's continuous wake-word recognizer hears 'Netra' in the open portal tab.")
add_numbered(doc, "Widget cancels TTS, starts a single-shot SpeechRecognition session, plays a brief 'Yes?' chirp.")
add_numbered(doc, "User says 'resolve I N C zero zero zero one two three four'.")
add_numbered(doc, "Browser converts speech to text via the Web Speech API.")
add_numbered(doc, "Widget POSTs { transcript, pending } to /api/<scope>/voice/command.")
add_numbered(doc, "Scripted REST handler invokes NetraIntent.parse() — regex matches 'resolve INC\\d+' and returns intent { action: 'resolve', args: { ticket_number: 'INC0001234' } }.")
add_numbered(doc, "NetraResponder.handle() dispatches to NetraTools.resolveTicket(), which sets state=6 via GlideRecord; NetraContext records the focused ticket so 'that' can refer to it.")
add_numbered(doc, "Responder composes a varied human-friendly reply, e.g. 'Ticket I N C zero zero zero one two three four is marked resolved.'")
add_numbered(doc, "Widget speaks the reply via SpeechSynthesis; ARIA live region announces it to screen-readers; the visible ticket list refreshes.")

page_break(doc)

# ============================================================
# 4. Component reference
# ============================================================
add_h1(doc, "4. Component reference")

add_h2(doc, "4.1 Script Includes (server-side)")
add_table(doc,
    headers=["Name", "Purpose", "Lines"],
    rows=[
        ["NetraIntent",     "Regex-based natural-language intent parser; handles 30+ intents including smalltalk, multi-turn continuations, ticket numbers, durations.", "~270"],
        ["NetraTools",      "GlideRecord operations on incidents, approvals, watchlist, and user prefs (pause/resume, voice mode).", "~200"],
        ["NetraResponder",  "Orchestrates every intent: calls NetraTools / NetraKnowledge / NetraChat / NetraContext / NetraNavigator / NetraSummarizer and composes varied spoken replies.", "~310"],
        ["NetraScanner",    "Periodic watcher invoked by the scheduled job; scans incidents / changes / sc_tasks / approvals for new items assigned to each user.", "~190"],
        ["NetraKnowledge",  "Keyword-scored search over the kb_knowledge table; returns ranked articles. No external API.", "~100"],
        ["NetraChat",       "Reads unread Connect Chat messages and posts replies to the user's most recent session. Gracefully no-ops if Connect is not enabled.", "~80"],
        ["NetraSummarizer", "Extractive summary using TextRank-lite (frequency-weighted sentence scoring). Used for long ticket descriptions and KB articles.", "~60"],
        ["NetraContext",    "Per-user session memory: focused ticket, last spoken utterance. Allows 'resolve that' / 'read it again' references.", "~70"],
        ["NetraNavigator",  "Maps spoken destinations ('open my profile', 'go to approvals') to portal URLs.", "~50"],
    ],
    col_widths_in=[1.5, 4.2, 0.8],
)

add_h2(doc, "4.2 Business Rule")
add_table(doc,
    headers=["Name", "Trigger", "Behaviour"],
    rows=[
        ["Netra Notify On Comment",
         "Async, after insert on sys_journal_field where name='incident' and element in (comments, work_notes)",
         "Resolves the incident's caller and assignee; if either is NOT the comment author, enqueues a notification for them with the spoken message and ticket number."],
    ],
    col_widths_in=[1.6, 2.4, 2.5],
)

add_h2(doc, "4.3 Scheduled Job")
add_table(doc,
    headers=["Name", "Frequency", "Script"],
    rows=[
        ["Netra Watch", "Every 3 minutes", "Delegates to NetraScanner.run() — iterates active users with prefs, scans assignments / approvals / catalog tasks since each user's last_scan_time, enqueues deduplicated notifications."],
    ],
    col_widths_in=[1.2, 1.4, 3.9],
)

add_h2(doc, "4.4 Scripted REST API")
add_p(doc, "Base path:  /api/<scope>/voice")
add_table(doc,
    headers=["Method", "Path", "Auth", "Purpose"],
    rows=[
        ["POST", "/command",       "Required",   "Accepts { transcript, pending }; returns { message, intent, data, pending, navigate, paused, ... }."],
        ["GET",  "/notifications", "Required",   "Returns and consumes the user's pending notifications; honours pause window."],
    ],
    col_widths_in=[0.8, 1.5, 0.8, 3.4],
)

add_h2(doc, "4.5 Service Portal widget")
add_image(doc, "widget-mockup", width_in=5.5, caption="Figure 5 — netra-mic widget, fixed bottom-right inside the Service Portal.")
add_p(doc, "Widget id: netra-mic")
add_bullet(doc, "Template — floating dock (mic, status, wake-toggle, pause / resume, help), pending-question banner, paused banner, conversation panel, sr-only ARIA live region.")
add_bullet(doc, "Client script — wake-word loop, command capture, TTS playback, notification polling every 8s, push-to-talk via Alt+N, Esc to stop.")
add_bullet(doc, "Server script — ensures one user-pref row exists for the current user (implicit opt-in on first load).")
add_bullet(doc, "Stylesheet — high-contrast accessible CSS; status-driven colours; pulsing red mic while listening.")

page_break(doc)

# ============================================================
# 5. Data model
# ============================================================
add_h1(doc, "5. Data model")
add_image(doc, "data-model", width_in=6.5, caption="Figure 3 — Four custom tables, all in the netra scope.")
add_p(doc, "Four custom tables, all created by the Fix Script under the netra scope.")

add_h3(doc, "5.1 <scope>_notification")
add_p(doc, "Queue of pending spoken announcements. The widget polls /notifications every 8 seconds; the endpoint marks each row delivered and the widget speaks it.")
add_table(doc,
    headers=["Column", "Type", "Notes"],
    rows=[
        ["user",          "Reference (sys_user)", "Recipient"],
        ["ticket_sys_id", "String(32)",           "Source record sys_id"],
        ["ticket_number", "String(32)",           "Display number, used for read-back"],
        ["kind",          "String(40)",           "comment | work_note | incident_assigned | change_assigned | sc_task_assigned | approval_requested"],
        ["message",       "String(1000)",         "Full spoken text"],
        ["delivered",     "True/False (default false)", "Set true once spoken"],
        ["delivered_at",  "Date/Time",            "Timestamp of consumption"],
    ],
    col_widths_in=[1.7, 1.7, 3.1],
)

add_h3(doc, "5.2 <scope>_user_pref")
add_p(doc, "One row per user, created on first widget load. Holds pause window and watch toggles.")
add_table(doc,
    headers=["Column", "Type", "Notes"],
    rows=[
        ["user",              "Reference (sys_user)",         "Unique"],
        ["active",            "True/False (default true)",    "Soft-disable per user"],
        ["paused_until",      "Date/Time",                    "Null = not paused"],
        ["last_scan_time",    "Date/Time",                    "Watermark used by NetraScanner"],
        ["watch_assignments", "True/False (default true)",    "Notify on new incident/change/task assignments"],
        ["watch_comments",    "True/False (default true)",    "Notify on new ticket comments"],
        ["watch_approvals",   "True/False (default true)",    "Notify on new pending approvals"],
        ["voice_mode",        "String(20) (default normal)",  "terse | normal | verbose"],
    ],
    col_widths_in=[1.7, 2.0, 2.8],
)

add_h3(doc, "5.3 <scope>_context")
add_p(doc, "Per-user conversational memory. Allows references like 'resolve that' or 'read it again'.")
add_table(doc,
    headers=["Column", "Type", "Notes"],
    rows=[
        ["user",           "Reference (sys_user)", "Unique"],
        ["focus_table",    "String(40)",           "e.g. incident, change_request"],
        ["focus_sys_id",   "String(32)",           "Last record discussed"],
        ["focus_number",   "String(32)",           "Display number for read-back"],
        ["focus_set_at",   "Date/Time",            "TTL — 12 hours"],
        ["last_utterance", "String(1000)",         "Last spoken reply (for 'repeat that')"],
    ],
    col_widths_in=[1.7, 1.7, 3.1],
)

add_h3(doc, "5.4 <scope>_watchlist")
add_p(doc, "User-curated set of records to receive proactive alerts on, beyond caller/assignee.")
add_table(doc,
    headers=["Column", "Type", "Notes"],
    rows=[
        ["user",          "Reference (sys_user)", "Owner"],
        ["record_table",  "String(40)",           "Watched table"],
        ["record_sys_id", "String(32)",           "Watched record"],
        ["record_number", "String(32)",           "Display number for spoken playback"],
    ],
    col_widths_in=[1.7, 1.7, 3.1],
)

page_break(doc)

# ============================================================
# 6. Voice interaction flow
# ============================================================
add_h1(doc, "6. Voice interaction flow")

add_h2(doc, "6.1 State machine (widget)")
add_code(doc, r"""
              [wake word OR Alt+N OR mic-click]
                            |
        +-----+    (idle)   v   (status:listening)
        |     +---------> [LISTENING]
        |                   |
        |     mic stream    | Web Speech STT emits final transcript
        |     ends          v
        |               [THINKING] -- POST /command -->
        |                   |
        |   pending=null    v        response
        +---<-- [SPEAKING] <----- TTS reply
                            |
                            +-- if pending=<state>, next utterance is
                                interpreted as the answer to that question
""")

add_h2(doc, "6.2 Multi-turn dialogue")
add_p(doc, "Three follow-up flows are supported:")
add_bullet(doc, "ticket_description — triggered by 'open a ticket' without details, asks 'What's the issue?'.")
add_bullet(doc, "resolve_number — triggered by 'resolve a ticket', asks 'Which I N C number?'.")
add_bullet(doc, "pause_duration — triggered by 'pause', asks 'For how many hours should I pause?'.")
add_p(doc, "The server returns pending='<state>' in its response; the widget echoes that back on the next user utterance, and NetraIntent.parse() routes accordingly.")

add_h2(doc, "6.3 Proactive interruption")
add_image(doc, "sequence-proactive", width_in=6.5, caption="Figure 4 — Sequence diagram for a proactive notification (assignment).")
add_p(doc, "Two parallel sources feed the notification queue:")
add_numbered(doc, "Business Rule on sys_journal_field — fires instantly when a colleague comments on a ticket where the user is caller or assignee. Sub-second latency.")
add_numbered(doc, "Scheduled Job 'Netra Watch' — runs every 3 minutes. NetraScanner iterates active users in the user_pref table and queries incident/change_request/sc_task/sysapproval_approver for new rows since the user's last_scan_time.")
add_p(doc, "Each enqueued row is deduplicated by (user, ticket_sys_id, kind). The widget polls every 8 seconds; on a hit, it speaks the message — unless the user has paused notifications (the endpoint returns empty during pause windows).")

page_break(doc)

# ============================================================
# 7. Voice command grammar
# ============================================================
add_h1(doc, "7. Voice command grammar")
add_p(doc, "All intents are recognised by regex patterns in NetraIntent.parse(). The matcher is permissive — multiple verb synonyms map to the same action.")

add_h2(doc, "7.1 Tickets")
add_table(doc,
    headers=["You say (any of)", "Intent", "Action"],
    rows=[
        ["create / open / log / raise / file / report a ticket for X", "create",       "Opens an incident, reads back the new INC number"],
        ["list / show my tickets, what's on my plate",                  "list",         "Reads up to 5 open tickets"],
        ["resolve / close / mark resolved INC0001234",                  "resolve",      "Sets state=6, close_code, close_notes"],
        ["update / comment on INC0001234 with X",                       "update",       "Adds X as a customer-visible comment"],
        ["status of / tell me about INC0001234",                        "status",       "Reads state, priority, assignee"],
        ["read / read me INC0001234",                                   "read",         "Full description + most recent comment, auto-summarised if long"],
        ["search tickets for X / find tickets about X",                 "search_tickets","Keyword search over short_description + description"],
        ["resolve that / read it again",                                "*_context",    "Acts on the focused ticket (see NetraContext)"],
    ],
    col_widths_in=[3.0, 1.0, 2.5],
)

add_h2(doc, "7.2 Approvals")
add_table(doc,
    headers=["You say", "Intent", "Action"],
    rows=[
        ["list / any pending approvals",            "list_approvals", "Reads all sysapproval_approver rows in 'requested' state"],
        ["approve CHG0001234 / approve RITM0010001", "approve",        "Sets state=approved with a comment"],
        ["reject / deny CHG0001234",                 "reject",         "Sets state=rejected"],
    ],
    col_widths_in=[2.5, 1.2, 2.8],
)

add_h2(doc, "7.3 Knowledge base")
add_table(doc,
    headers=["You say", "Intent", "Action"],
    rows=[
        ["search knowledge for X / how do X",     "kb_search", "Keyword-scored search over kb_knowledge; reads top 3 titles"],
        ["read KB0010001",                         "kb_read",   "Reads the article body, summarised if long"],
    ],
    col_widths_in=[2.8, 1.0, 2.7],
)

add_h2(doc, "7.4 Pause / resume / mode")
add_table(doc,
    headers=["You say", "Intent", "Action"],
    rows=[
        ["pause (then 'two hours')",      "pause_ask / pause_for", "Stores paused_until on user_pref; widget honours immediately"],
        ["resume / wake up / come back",  "resume",                 "Clears paused_until"],
        ["switch to terse / verbose mode","voice_mode",             "Stores voice_mode preference (terse / normal / verbose)"],
    ],
    col_widths_in=[2.8, 1.4, 2.3],
)

add_h2(doc, "7.5 Smalltalk and meta")
add_table(doc,
    headers=["You say", "Reply style"],
    rows=[
        ["hi / hello / good morning",        "Time-of-day greeting"],
        ["thanks / thank you / good job",    "Varied acknowledgment"],
        ["how are you / who are you",        "Short smalltalk"],
        ["repeat that / say it again",       "Re-reads the last spoken reply"],
        ["stop / cancel / quiet",            "Interrupts speech, clears pending state"],
        ["help / what can you do",           "Lists the available verbs"],
    ],
    col_widths_in=[3.0, 3.5],
)

page_break(doc)

# ============================================================
# 8. Accessibility design
# ============================================================
add_h1(doc, "8. Accessibility design")
add_bullet(doc, "ARIA live region (role=status, aria-live=assertive) announces every state change so JAWS, NVDA, and VoiceOver speak it.")
add_bullet(doc, "Every button has an explicit aria-label; the mic button uses aria-pressed and aria-busy to indicate recording / processing state.")
add_bullet(doc, "Keyboard reachable: Alt+N push-to-talk, Esc cancels speech, all buttons in focus order.")
add_bullet(doc, "High contrast: white background, deep blue (#1A56DB / #0E2E7A) primary, deep red (#C81E1E) for recording — passes WCAG 2.1 AA contrast checks at 4.5:1 or higher.")
add_bullet(doc, "Pulsing mic ring during recording — visual reinforcement for low-vision users.")
add_bullet(doc, "Ticket numbers always read letter-by-digit ('I N C zero zero zero one two three four') so the user can write them down.")
add_bullet(doc, "Voice mode preference (terse / normal / verbose) lets the user tune reply length.")
add_bullet(doc, "Optional text-only mode (typed input box) for environments where mic use is impractical.")
add_bullet(doc, "Skip link at the top of the widget for quick keyboard access.")

page_break(doc)

# ============================================================
# 9. Security & data residency
# ============================================================
add_h1(doc, "9. Security & data residency")

add_h2(doc, "9.1 Authentication & authorisation")
add_bullet(doc, "Scripted REST endpoints require authentication; ACL checks fall through to the scoped app's defaults.")
add_bullet(doc, "Every GlideRecord operation in NetraTools is scoped to caller_id = current user (or, for read, caller_id OR assigned_to). Users cannot see or modify tickets they don't own.")
add_bullet(doc, "Approval decisions are gated by the standard sysapproval_approver record (approver = current user, state = requested).")
add_bullet(doc, "Knowledge searches respect the platform's KB ACLs (read-only).")
add_bullet(doc, "The scoped application means uninstall (delete app) removes every record automatically.")

add_h2(doc, "9.2 Data residency")
add_table(doc,
    headers=["Data class", "Where it lives", "External?"],
    rows=[
        ["Microphone audio",        "Browser → Web Speech API server (Google for Chrome, Microsoft for Edge, Apple for Safari) for STT only.", "Yes — to browser STT provider; never to a third-party we control"],
        ["Transcribed text",        "Sent from browser to ServiceNow Scripted REST API, stored in journal/ticket fields as the user dictates.", "No"],
        ["Ticket / KB / chat data", "ServiceNow tables only.", "No"],
        ["TTS playback",            "Browser speech synthesis; reply text sent to browser TTS engine.", "Audio is rendered locally; text was already in-instance"],
        ["Notifications queue",     "Custom <scope>_notification table.", "No"],
        ["User preferences",        "Custom <scope>_user_pref table.", "No"],
        ["Session context",         "Custom <scope>_context table.", "No"],
    ],
    col_widths_in=[1.6, 2.6, 2.3],
)
add_p(doc, "For users who require strict in-instance processing (no audio leaving the browser host), the text-only mode bypasses Web Speech entirely — the user types commands instead.", italic=True)

add_h2(doc, "9.3 Audit trail")
add_bullet(doc, "Every ticket mutation goes through standard GlideRecord — sys_audit / sys_history track who changed what.")
add_bullet(doc, "NetraScanner logs scan counts to System Logs → Script Log Statements (filter by [NetraScanner]).")
add_bullet(doc, "The notification queue retains both message and delivered_at for replay / debugging.")

page_break(doc)

# ============================================================
# 10. Installation guide
# ============================================================
add_h1(doc, "10. Installation guide")

add_h2(doc, "10.1 Prerequisites")
add_bullet(doc, "A ServiceNow developer instance (or any non-prod instance).")
add_bullet(doc, "Admin credentials on the target instance.")
add_bullet(doc, "Chrome or Edge browser for end users (Web Speech API support).")

add_h2(doc, "10.2 Step-by-step")
add_numbered(doc, "Create the scoped application. In ServiceNow, navigate to System Applications → My Company Applications → Create new → Start from scratch. Enter Name = 'netra'. ServiceNow auto-generates the scope as x_<companyId>_netra.")
add_numbered(doc, "Download the deployable Update Set: netra-snow/update-set/netra-deployable.xml.")
add_numbered(doc, "In ServiceNow, navigate to System Update Sets → Retrieved Update Sets → Import Update Set from XML. Upload the file.")
add_numbered(doc, "Open the loaded update set → Preview Update Set → resolve any (rare) collisions → Commit Update Set.")
add_numbered(doc, "Navigate to System Definition → Fix Scripts. Open 'Netra Install' and click Run Fix Script. Watch the script output — it will print every table, Script Include, Business Rule, Scheduled Job, REST resource and Widget it creates, then end with '=== Netra install complete ==='.")
add_numbered(doc, "Add the widget to a portal page: Service Portal → Service Portal Configuration → Designer → open the home page → drag 'Netra Mic' into any container → Save.")
add_numbered(doc, "Open /sp in Chrome or Edge. Allow microphone when prompted. Say 'Netra' or click the floating mic to begin.")

add_h2(doc, "10.3 Verification checklist")
add_bullet(doc, "Inside the netra app in Studio, the Application Files panel should show 9 Script Includes, 4 Tables, 1 Business Rule, 1 Scheduled Job, 1 Scripted REST Service with 2 Resources, and 1 Widget.")
add_bullet(doc, "System Definition → Scheduled Jobs → filter by Application = Netra. 'Netra Watch' should be Active with Next action within 3 minutes.")
add_bullet(doc, "Visit the widget URL — the floating dock should appear; say 'Netra' and a 'Yes?' chirp should play.")
add_bullet(doc, "Speak 'create a ticket for my email is broken' — within a few seconds a new INC should appear in your ticket list and Netra should read back the number.")

add_h2(doc, "10.4 Uninstall")
add_p(doc, "Navigate to System Applications → All Available Applications → My Apps → Netra → Delete. ServiceNow removes every record in the scope automatically.")

page_break(doc)

# ============================================================
# 11. Operational concerns
# ============================================================
add_h1(doc, "11. Operational concerns")

add_h2(doc, "11.1 Performance")
add_bullet(doc, "Scheduled job runs in <500 ms per user on a typical dev instance (10–20 users with a handful of open records each).")
add_bullet(doc, "Each /command round-trip averages 80–150 ms server-side (regex + 1–3 GlideRecord queries).")
add_bullet(doc, "Notification polling: GET /notifications is a single indexed query; widget polls every 8 s. No measurable impact on instance load.")

add_h2(doc, "11.2 Failure handling")
add_bullet(doc, "Wake-word recognition restarts automatically on 'no-speech', 'aborted', and stream-end events.")
add_bullet(doc, "If the Scripted REST call fails, the widget falls back to a generic apology and resets to idle — never gets stuck.")
add_bullet(doc, "The scheduled job's run loop wraps each user in try/catch so a bad row never aborts the whole scan.")
add_bullet(doc, "If Connect Chat is not enabled on the instance, NetraChat reports the feature as disabled rather than throwing.")

add_h2(doc, "11.3 Observability")
add_bullet(doc, "Server-side log lines tag with [NetraScanner] and [NetraResponder] for quick grep in System Logs.")
add_bullet(doc, "Widget conversation history is visible in the panel for the current session — useful for support.")
add_bullet(doc, "The notification table retains every delivered message; review it for replay.")

page_break(doc)

# ============================================================
# 12. Test plan
# ============================================================
add_h1(doc, "12. Test plan")

add_h2(doc, "12.1 Functional test cases")
add_table(doc,
    headers=["ID", "Scenario", "Expected outcome"],
    rows=[
        ["T-01", "Speak 'create a ticket for VPN issues'", "New INC appears; Netra reads back its number"],
        ["T-02", "Speak 'list my tickets' with 3 open tickets", "Netra enumerates them in order of last update"],
        ["T-03", "Speak 'resolve I N C zero zero zero one two three four'", "Ticket state changes to Resolved; visual update in list"],
        ["T-04", "Speak 'update INC0001234 with I rebooted the laptop'", "Comment appears on the ticket; visible to assignee"],
        ["T-05", "Speak 'pause' then 'two hours'", "User_pref.paused_until set 2 h ahead; purple banner shows; notifications suppressed"],
        ["T-06", "Speak 'resume' while paused", "Paused_until cleared; banner disappears; notifications flow"],
        ["T-07", "Speak 'help'", "Netra reads the list of supported commands"],
        ["T-08", "Have an admin user comment on the user's ticket", "Within 8 s the widget speaks the comment text and author"],
        ["T-09", "Have an admin user assign an incident to the user", "Within 3 min (next scheduler run) Netra announces the assignment"],
        ["T-10", "Speak 'search knowledge for password reset'", "Netra reads top 3 article titles + numbers"],
        ["T-11", "Speak 'read KB0010001'", "Netra reads the article body, summarised to ~4 sentences"],
        ["T-12", "Speak 'open my approvals' with one pending approval", "Netra reads the approval's subject"],
        ["T-13", "Speak 'approve CHG0001234'", "Approval state changes to approved"],
        ["T-14", "Speak 'watch INC0001234' then 'list my watchlist'", "Ticket is added; later list call includes it"],
        ["T-15", "Speak 'open my tickets'", "Browser navigates to the ticket list page"],
        ["T-16", "Press Alt+N", "Recording starts even without saying the wake word"],
        ["T-17", "Speak 'stop' while Netra is talking", "TTS cancels immediately"],
        ["T-18", "Speak 'repeat that'", "Netra re-reads her last reply"],
    ],
    col_widths_in=[0.6, 3.4, 2.5],
)

add_h2(doc, "12.2 Accessibility test cases")
add_bullet(doc, "Operate every function with keyboard only — Tab through controls, activate with Enter / Space.")
add_bullet(doc, "Verify NVDA / JAWS / VoiceOver announce every state change.")
add_bullet(doc, "Audit colour contrast with a WCAG checker — must pass AA at minimum.")
add_bullet(doc, "Run high-contrast mode in OS settings; verify Netra remains legible.")

page_break(doc)

# ============================================================
# 13. Roadmap
# ============================================================
add_h1(doc, "13. Roadmap & future work")

add_h2(doc, "13.1 Near-term (next minor)")
add_bullet(doc, "Voice forms — fill out catalog item variables by voice ('request a new laptop').")
add_bullet(doc, "Confirmation prompts for destructive verbs ('Are you sure you want to resolve INC0001234?').")
add_bullet(doc, "Multi-table search (changes, problems, requests) — currently only incidents.")

add_h2(doc, "13.2 Mid-term")
add_bullet(doc, "ServiceNow Virtual Agent fallback — when intent parsing returns 'unknown', forward to VA for richer NLU.")
add_bullet(doc, "Custom wake-word training via a small client-side ML model.")
add_bullet(doc, "Localised wake words and intent patterns for additional languages.")

add_h2(doc, "13.3 Long-term")
add_bullet(doc, "Mobile-optimised Now Mobile shell.")
add_bullet(doc, "Voice-driven runbook execution (Flow Designer integration).")
add_bullet(doc, "Optional integration with Now Assist for users on Pro / Enterprise plans (kept opt-in, never default, to preserve the no-external-API stance).")

page_break(doc)

# ============================================================
# Appendix A — source file index
# ============================================================
add_h1(doc, "Appendix A — Source file index")
add_p(doc, "Every artifact lives in the netra-snow/ folder of the repository.")
add_table(doc,
    headers=["Path", "Purpose"],
    rows=[
        ["update-set/netra-deployable.xml",                     "The single-file Update Set you import (~150 KB)"],
        ["install/setup-netra.js",                              "Alternative: pasteable Background Script (~103 KB)"],
        ["preview/index.html",                                  "Self-contained browser preview with mocked backend"],
        ["source/fix_script/netra-install.js",                  "The Fix Script template embedded into the Update Set"],
        ["source/script_includes/NetraIntent.js",               "Regex intent parser"],
        ["source/script_includes/NetraTools.js",                "GlideRecord operations"],
        ["source/script_includes/NetraResponder.js",            "Spoken reply composer"],
        ["source/script_includes/NetraScanner.js",              "Periodic assignment/approval watcher"],
        ["source/script_includes/NetraKnowledge.js",            "Knowledge base search"],
        ["source/script_includes/NetraChat.js",                 "Connect Chat reader/writer"],
        ["source/script_includes/NetraSummarizer.js",           "Extractive summary"],
        ["source/script_includes/NetraContext.js",              "Per-user session memory"],
        ["source/script_includes/NetraNavigator.js",            "Portal URL routing"],
        ["source/scripted_rest/command.js",                     "POST /api/<scope>/voice/command"],
        ["source/scripted_rest/notifications.js",               "GET  /api/<scope>/voice/notifications"],
        ["source/business_rule/netra_notify_on_comment.js",     "Async after-insert on sys_journal_field"],
        ["source/scheduled_jobs/netra_watch.js",                "Wraps NetraScanner.run()"],
        ["source/widget/*",                                     "Service Portal widget (template / client / server / SCSS)"],
        ["scripts/build-deployable.ps1",                        "Regenerates the Update Set XML"],
        ["scripts/build-setup-script.ps1",                      "Regenerates the Background Script"],
        ["scripts/build-design-doc.py",                         "Regenerates this document"],
        ["scripts/serve-preview.ps1",                           "Local HTTP server for the preview"],
    ],
    col_widths_in=[3.2, 3.3],
)

# Footer (page numbers) — leave default

# Save
doc.save(OUT)
print(f"Wrote {OUT}")
print(f"  size: {os.path.getsize(OUT):,} bytes")
